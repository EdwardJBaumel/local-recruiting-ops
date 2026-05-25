"""
Local resume reference store.

Everything lives under data/resume/ on the user's machine. Nothing leaves the
box unless the user forwards it to the match pipeline (which also runs locally).

Layout:
  data/resume/
    source.pdf | source.docx     - original uploaded file
    parsed.txt                    - extracted text
    additional_notes.txt          - user's free-text notes
    metadata.json                 - filename, upload time, char count
"""
from __future__ import annotations

import base64
import datetime as _dt
import json
import logging
import threading
from pathlib import Path

logger = logging.getLogger("lantern.resume_store")

_ACCEPTED_EXT = {".pdf", ".docx", ".txt", ".md"}
_MAX_BYTES = 10 * 1024 * 1024  # 10 MB; a real resume is <1 MB

# Serialise all writes so concurrent uploads/notes-saves don't interleave.
_lock = threading.Lock()


def _store_dir(data_dir: Path) -> Path:
    d = data_dir / "resume"
    d.mkdir(parents=True, exist_ok=True)
    return d


# Crash-safe writes live in core.io_safe now so there's one copy. These
# local aliases keep the original callsites intact and make the intent
# obvious at the point of use.
from core.io_safe import write_bytes_atomic as _atomic_write_bytes
from core.io_safe import write_text_atomic as _atomic_write_text  # noqa: F401


# ──────────────────────────────────────────────────────────────────
# Parsing
# ──────────────────────────────────────────────────────────────────
def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError(
            "pypdf not installed. Run: pip install pypdf"
        )
    reader = PdfReader(str(path))
    out = []
    for page in reader.pages:
        try:
            out.append(page.extract_text() or "")
        except Exception as e:
            logger.warning("PDF page extract failed: %s", e)
    text = "\n".join(out).strip()
    if not text:
        raise RuntimeError(
            "Could not extract any text from the PDF. If it is a scan, "
            "save it as a text-based PDF (re-export from Word/Google Docs) "
            "and try again."
        )
    return text


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx not installed. Run: pip install python-docx"
        )
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull plain text out of tables (common in resumes).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    text = "\n".join(paragraphs).strip()
    if not text:
        raise RuntimeError("DOCX parsed OK but contained no text.")
    return text


def _parse_plain(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace").strip()


def _parse_by_ext(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext == ".docx":
        return _parse_docx(path)
    if ext in (".txt", ".md"):
        return _parse_plain(path)
    raise RuntimeError(f"Unsupported file type: {ext}")


# ──────────────────────────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────────────────────────
def save_upload(data_dir: Path, filename: str, content_b64: str) -> dict:
    """Decode a base64 upload, save the original, parse to text.
    Raises ValueError/RuntimeError on bad input.
    """
    ext = Path(filename).suffix.lower()
    if ext not in _ACCEPTED_EXT:
        raise ValueError(
            f"Unsupported file type '{ext}'. Accepted: {sorted(_ACCEPTED_EXT)}"
        )

    try:
        # Strip a data URL prefix if the browser included one.
        if "," in content_b64[:80] and content_b64.lstrip().startswith("data:"):
            content_b64 = content_b64.split(",", 1)[1]
        raw = base64.b64decode(content_b64, validate=True)
    except Exception as e:
        raise ValueError(f"content_base64 is not valid base64: {e}")

    if len(raw) == 0:
        raise ValueError("Uploaded file is empty.")
    if len(raw) > _MAX_BYTES:
        raise ValueError(
            f"Uploaded file is {len(raw)/1_048_576:.1f} MB, max is "
            f"{_MAX_BYTES/1_048_576:.0f} MB."
        )

    with _lock:
        d = _store_dir(data_dir)

        # Clear any existing original first so we don't leave a .pdf AND a .docx.
        for old in d.glob("source.*"):
            try:
                old.unlink()
            except OSError:
                pass

        src_path = d / f"source{ext}"
        _atomic_write_bytes(src_path, raw)

        try:
            text = _parse_by_ext(src_path)
        except Exception:
            # Leave the original in place so the user can see what they uploaded,
            # but don't pretend we parsed it.
            raise

        _atomic_write_text(d / "parsed.txt", text)

        meta = {
            "filename": filename,
            "size_bytes": len(raw),
            "char_count": len(text),
            "uploaded_at": _dt.datetime.now(_dt.timezone.utc).isoformat(timespec="seconds"),
        }
        _atomic_write_text(d / "metadata.json", json.dumps(meta, indent=2))

    logger.info("Resume stored: %s (%d chars extracted)", filename, len(text))
    return meta


def save_notes(data_dir: Path, notes: str) -> dict:
    if not isinstance(notes, str):
        raise ValueError("notes must be a string")
    with _lock:
        d = _store_dir(data_dir)
        _atomic_write_text(d / "additional_notes.txt", notes)
    return {"char_count": len(notes)}


def read_current(data_dir: Path) -> dict:
    """Returns the current resume state as a JSON-safe dict."""
    d = _store_dir(data_dir)
    meta_path = d / "metadata.json"
    parsed_path = d / "parsed.txt"
    notes_path = d / "additional_notes.txt"

    meta = {}
    if meta_path.exists():
        try:
            meta = json.loads(meta_path.read_text())
        except Exception:
            meta = {}

    parsed_text = parsed_path.read_text(encoding="utf-8") if parsed_path.exists() else ""
    notes = notes_path.read_text(encoding="utf-8") if notes_path.exists() else ""

    return {
        "has_resume": parsed_path.exists(),
        "metadata": meta,
        "parsed_text": parsed_text,
        "additional_notes": notes,
    }


def clear(data_dir: Path) -> None:
    with _lock:
        d = _store_dir(data_dir)
        for p in d.glob("*"):
            try:
                p.unlink()
            except OSError:
                pass


def get_profile_text(data_dir: Path) -> str | None:
    """Combined text for the match pipeline. None if no resume uploaded.
    Returns resume text with additional_notes appended if present.
    """
    state = read_current(data_dir)
    if not state["has_resume"]:
        return None
    parts = [state["parsed_text"]]
    if state["additional_notes"].strip():
        parts.append("\nAdditional notes from the candidate:\n" + state["additional_notes"])
    return "\n\n".join(parts)
